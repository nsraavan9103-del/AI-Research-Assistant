"""
Script to generate a comprehensive Word document from markdown documentation files
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

def set_cell_background(cell, fill_color):
    """Set cell background color"""
    tcPr = cell._element.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:shd')
    tcVAlign.set(qn('w:fill'), fill_color)
    tcPr.append(tcVAlign)

def add_heading_with_style(doc, text, level):
    """Add a heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph_with_style(doc, text, is_code=False, indent=0):
    """Add a paragraph with proper styling"""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Inches(indent * 0.25)
    
    if is_code:
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    return para

def parse_markdown_table(lines, start_idx):
    """Parse markdown table and return table data and next line index"""
    if start_idx + 2 >= len(lines):
        return None, start_idx
    
    header_line = lines[start_idx].strip()
    separator_line = lines[start_idx + 1].strip()
    
    if '|' not in header_line or '|' not in separator_line:
        return None, start_idx
    
    headers = [h.strip() for h in header_line.split('|')]
    headers = [h for h in headers if h]
    
    rows = []
    idx = start_idx + 2
    
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or '|' not in line:
            break
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c]
        if len(cells) == len(headers):
            rows.append(cells)
        idx += 1
    
    return (headers, rows), idx

def generate_docx():
    """Generate comprehensive Word document from markdown files"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('AI Research Assistant')
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Complete Technical Documentation')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(51, 102, 153)
    
    doc.add_paragraph()
    
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_text = "Version: 1.0 | Last Updated: March 2026\nTech Stack: FastAPI • React • PostgreSQL • LangChain • Ollama"
    run = meta.add_run(meta_text)
    run.font.size = Pt(11)
    run.font.italic = True
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. Software Modules',
        '3. System Architecture',
        '4. Data Flow Diagram',
        '5. Database Design',
        '6. API Endpoints',
        '7. Frontend Components',
        '8. RAG Pipeline Details',
        '9. Setup & Installation',
        '10. Usage Guide',
        '11. Performance Metrics',
        '12. Troubleshooting'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Read comprehensive documentation
    doc_path = r'c:\Users\Sraavan\Desktop\React and Python\Main-Project\AI-Research-Assistant\Documentation\COMPREHENSIVE_PROJECT_DOCUMENTATION.md'
    
    with open(doc_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into sections for processing
    lines = content.split('\n')
    
    current_section = []
    for i, line in enumerate(lines):
        current_section.append(line)
        
        # Process sections
        if line.startswith('## '):
            if len(current_section) > 1:
                process_section(doc, current_section[:-1])
            current_section = [line]
        
        # Add page breaks after major sections
        if line.startswith('## ') and i > 0:
            pass
    
    # Process remaining section
    if current_section:
        process_section(doc, current_section)
    
    # Appendices
    doc.add_page_break()
    doc.add_heading('Appendix A: Quick Reference', level=1)
    
    doc.add_heading('Technology Stack', level=2)
    tech_items = [
        'Backend: FastAPI + Python 3.10+',
        'Frontend: React 18 + Vite + TypeScript',
        'Database: PostgreSQL (or SQLite for dev)',
        'Vector DB: FAISS (local)',
        'LLM: Ollama (phi3, 7B)',
        'Embeddings: nomic-embed-text',
        'Reranker: BGE-Reranker-v2-m3',
        'Task Queue: Celery + Redis',
        'Styling: TailwindCSS'
    ]
    for item in tech_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    doc.add_heading('Appendix B: File Structure', level=1)
    
    structure = """AI-Research-Assistant/
├── Backend/
│   ├── main.py
│   ├── models.py
│   ├── database.py
│   ├── auth.py
│   ├── rag.py
│   ├── requirements.txt
│   ├── api/
│   │   ├── routes/
│   │   └── dependencies/
│   ├── services/
│   │   ├── rag_chain.py
│   │   ├── retrieval/
│   │   └── cache.py
│   └── tasks/
├── frontend/
│   ├── src/
│   │   ├── components/
│   │   ├── pages/
│   │   ├── hooks/
│   │   ├── store/
│   │   └── App.jsx
│   ├── package.json
│   └── vite.config.js
├── Documentation/
│   ├── COMPREHENSIVE_PROJECT_DOCUMENTATION.md
│   ├── QUICK_START_GUIDE.md
│   ├── API_REFERENCE.md
│   └── ARCHITECTURE_AND_DESIGN.md
└── README.md"""
    
    doc.add_paragraph(structure, style='No Spacing')
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_para.add_run('\n\n--- End of Documentation ---\n\n')
    run.font.italic = True
    
    version_para = doc.add_paragraph()
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version_para.add_run('AI Research Assistant v1.0 | March 2026')
    run.font.size = Pt(9)
    run.font.italic = True
    
    # Save document
    output_path = r'c:\Users\Sraavan\Desktop\React and Python\Main-Project\AI-Research-Assistant\Documentation\AI_Research_Assistant_Documentation.docx'
    doc.save(output_path)
    
    print(f"✓ Word document created successfully!")
    print(f"📄 Location: {output_path}")
    print(f"📊 File size: {os.path.getsize(output_path) / 1024:.1f} KB")


def process_section(doc, lines):
    """Process a section of markdown and add to document"""
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Headers
        if line.startswith('#### '):
            add_heading_with_style(doc, line[5:].strip(), 4)
        elif line.startswith('### '):
            add_heading_with_style(doc, line[4:].strip(), 3)
        elif line.startswith('## '):
            add_heading_with_style(doc, line[3:].strip(), 2)
        elif line.startswith('# '):
            add_heading_with_style(doc, line[2:].strip(), 1)
        
        # Code blocks
        elif line.strip().startswith('```'):
            code_block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_block.append(lines[i])
                i += 1
            if code_block:
                code_text = '\n'.join(code_block).strip()
                para = doc.add_paragraph(code_text)
                para.style = 'Normal'
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                # Add background color to code paragraphs
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E8E8E8')
                para._element.get_or_add_pPr().append(shading_elm)
        
        # Tables (simple detection)
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_data, next_i = parse_markdown_table(lines, i)
            if table_data:
                headers, rows = table_data
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Header row
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                    # Style header
                    for run in hdr_cells[j].paragraphs[0].runs:
                        run.font.bold = True
                    set_cell_background(hdr_cells[j], 'D3D3D3')
                
                # Data rows
                for row in rows:
                    row_cells = table.add_row().cells
                    for j, cell in enumerate(row):
                        row_cells[j].text = cell
                
                i = next_i - 1
        
        # Bullet lists
        elif line.strip().startswith('- '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^\d+\.\s(.*)$', line.strip())
            if match:
                doc.add_paragraph(match.group(1), style='List Number')
        
        # Paragraphs
        elif line.strip() and not line.startswith('|'):
            add_paragraph_with_style(doc, line.strip())
        
        # Empty lines
        elif not line.strip():
            doc.add_paragraph()
        
        i += 1


if __name__ == '__main__':
    print("🔄 Generating Word document from markdown documentation...")
    print("📝 This may take a moment...\n")
    
    try:
        generate_docx()
        print("\n✨ Success! Your documentation is ready.")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
