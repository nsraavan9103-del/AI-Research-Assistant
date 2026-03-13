# ============================================================
# AI Research Assistant - Fixed Backend (main.py)
# Compatible with: Windows 10, Intel i5-10th Gen, GTX 1650Ti, 8GB RAM
# Fixed by: Claude (Anthropic) - March 2026
# ============================================================
import hashlib
import os
import uuid
import shutil
from pathlib import Path
from datetime import datetime, timedelta
from typing import Optional, List
from pydantic import BaseModel, Field

# ── Environment ──────────────────────────────────────────────
from dotenv import load_dotenv
load_dotenv()

# ── FastAPI ───────────────────────────────────────────────────
from fastapi import (
    FastAPI, Depends, HTTPException, UploadFile, File,
    status, BackgroundTasks
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# ── Auth / JWT ────────────────────────────────────────────────
from jose import JWTError, jwt
from passlib.context import CryptContext

# ── Database ──────────────────────────────────────────────────
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, Boolean
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session

# ── LangChain — CORRECTED IMPORTS (v0.2+) ────────────────────
from langchain_ollama import OllamaEmbeddings, OllamaLLM          # FIX: was langchain.embeddings / langchain.llms
from langchain_community.vectorstores import FAISS                 # FIX: was langchain.vectorstores
from langchain_community.document_loaders import PyPDFLoader       # FIX: was langchain.document_loaders
from langchain_text_splitters import RecursiveCharacterTextSplitter # FIX: was langchain.text_splitter
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import httpx

# ─────────────────────────────────────────────────────────────
# CONFIGURATION (all from .env — never hardcode secrets)
# ─────────────────────────────────────────────────────────────
SECRET_KEY = os.getenv("SECRET_KEY", "CHANGE-THIS-IN-PRODUCTION-USE-A-REAL-SECRET")
ALGORITHM = os.getenv("ALGORITHM", "HS256")
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "30"))
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:3b")           # Primary
OLLAMA_MODEL_FALLBACK = os.getenv("OLLAMA_MODEL_FALLBACK", "phi3:mini")  # Backup
EMBED_MODEL = os.getenv("EMBED_MODEL", "nomic-embed-text")


# ─────────────────────────────────────────────────────────────
# WINDOWS-COMPATIBLE PATHS (FIX: was Linux-style string paths)
# ─────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# ─────────────────────────────────────────────────────────────
# DATABASE SETUP
# ─────────────────────────────────────────────────────────────
SQLALCHEMY_DATABASE_URL = f"sqlite:///{BASE_DIR / 'research_app.db'}"
engine = create_engine(
    SQLALCHEMY_DATABASE_URL,
    connect_args={"check_same_thread": False}  # Required for SQLite + FastAPI
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True)
    email = Column(String, unique=True, index=True)
    hashed_password = Column(String)
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime, default=datetime.utcnow)


class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String)
    file_path = Column(String)
    owner_id = Column(Integer)
    session_id = Column(String, index=True)
    uploaded_at = Column(DateTime, default=datetime.utcnow)
    summary = Column(Text, nullable=True)


class ChatHistory(Base):
    __tablename__ = "chat_history"
    id = Column(Integer, primary_key=True, index=True)
    session_id = Column(String, index=True)
    question = Column(Text)
    answer = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)


Base.metadata.create_all(bind=engine)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ─────────────────────────────────────────────────────────────
# AUTH SETUP
# ─────────────────────────────────────────────────────────────
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

def _pre_hash(password: str) -> str:
    """
    Hashes the password with SHA-256 to ensure the length never exceeds 
    bcrypt's 72-byte limit, regardless of the user's input size.
    """
    return hashlib.sha256(password.encode('utf-8')).hexdigest()

def verify_password(plain: str, hashed: str) -> bool:
    # Pre-hash the plain text before verifying against the stored bcrypt hash
    pre_hashed = _pre_hash(plain)
    return pwd_context.verify(pre_hashed, hashed)

def get_password_hash(password: str) -> str:
    # Pre-hash the password to a fixed length before passing to bcrypt
    pre_hashed = _pre_hash(password)
    return pwd_context.hash(pre_hashed)


def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=15))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)


def get_user(db: Session, username: str):
    return db.query(User).filter(User.username == username).first()


def authenticate_user(db: Session, username: str, password: str):
    user = get_user(db, username)
    if not user or not verify_password(password, user.hashed_password):
        return None
    return user


async def get_current_user(
    token: str = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = get_user(db, username)
    if user is None:
        raise credentials_exception
    return user


# ─────────────────────────────────────────────────────────────
# VECTOR STORE CACHE — FIX: rebuild only when needed (not every request)
# ─────────────────────────────────────────────────────────────
_vectorstore_cache: dict = {}  # key: session_id → FAISS vectorstore


def get_vectorstore(session_id: str):
    return _vectorstore_cache.get(session_id)


def _load_documents_from_file(file_path: Path):
    """Return a list of langchain Document objects from any supported file type."""
    from langchain_community.document_loaders import PyPDFLoader, TextLoader
    from langchain_core.documents import Document as LCDoc

    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return PyPDFLoader(str(file_path)).load()
    elif ext == ".docx":
        try:
            import docx2txt
            text = docx2txt.process(str(file_path))
        except ImportError:
            from docx import Document as DocxDoc
            doc = DocxDoc(str(file_path))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [LCDoc(page_content=text or "", metadata={"source": file_path.name})]
    elif ext in (".xlsx", ".xls"):
        import openpyxl
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        docs = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append("\t".join(cells))
            sheet_text = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
            docs.append(LCDoc(page_content=sheet_text, metadata={"source": file_path.name, "sheet": sheet_name}))
        wb.close()
        return docs or [LCDoc(page_content="", metadata={"source": file_path.name})]
    else:
        # .txt and anything else — plain text
        return TextLoader(str(file_path), encoding="utf-8").load()


def build_vectorstore(session_id: str, file_path: Path):
    """Load document (any supported type), split text, build FAISS index, cache it."""
    documents = _load_documents_from_file(file_path)

    # FIX: chunk_size=500 is safe for phi3's 4096 token context window
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        length_function=len,
    )
    chunks = splitter.split_documents(documents)

    try:
        embeddings = OllamaEmbeddings(model=EMBED_MODEL)
        vectorstore = FAISS.from_documents(chunks, embeddings)
        _vectorstore_cache[session_id] = vectorstore
        return vectorstore, len(chunks)
    except httpx.ConnectError:
        raise HTTPException(
            status_code=503,
            detail="Ollama is not running. Please start it with: ollama serve"
        )



# ─────────────────────────────────────────────────────────────
# OLLAMA HEALTH CHECK
# ─────────────────────────────────────────────────────────────
def check_ollama():
    """Verify Ollama is reachable before trying inference."""
    try:
        import httpx
        r = httpx.get("http://localhost:11434/api/tags", timeout=3.0)
        return r.status_code == 200
    except Exception:
        return False


# ─────────────────────────────────────────────────────────────
# FASTAPI APP
# ─────────────────────────────────────────────────────────────
app = FastAPI(
    title="AI Research Assistant",
    description="Local RAG-powered research assistant using Ollama + FAISS",
    version="1.0.0"
)

# FIX: CORS — allows React on port 3000 to talk to FastAPI on port 8000
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000", 
        "http://localhost:5173",    # Added Vite default port
        "http://127.0.0.1:5173"     # Added local loopback alternative
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─────────────────────────────────────────────────────────────
# PYDANTIC SCHEMAS
# ─────────────────────────────────────────────────────────────
class UserCreate(BaseModel):
    username: str = Field(..., min_length=3, max_length=50)
    email: str
    password: str = Field(..., min_length=8, max_length=64)


class Token(BaseModel):
    access_token: str
    token_type: str


class QuestionRequest(BaseModel):
    question: str
    session_id: str


class AnswerResponse(BaseModel):
    answer: str
    session_id: str
    sources: List[str] = []


# ─────────────────────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────────────────────

@app.get("/")
def root():
    ollama_ok = check_ollama()
    return {
        "message": "AI Research Assistant is running",
        "ollama_status": "online" if ollama_ok else "OFFLINE — run: ollama serve",
        "models": {"llm": OLLAMA_MODEL, "embed": EMBED_MODEL}
    }


@app.get("/health")
def health():
    return {"status": "ok", "ollama": check_ollama()}


# ── Authentication ────────────────────────────────────────────

@app.post("/register", response_model=dict)
def register(user_data: UserCreate, db: Session = Depends(get_db)):
    if get_user(db, user_data.username):
        raise HTTPException(status_code=400, detail="Username already registered")
    hashed = get_password_hash(user_data.password)
    user = User(
        username=user_data.username,
        email=user_data.email,
        hashed_password=hashed
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return {"message": "User registered successfully", "username": user.username}


@app.post("/token", response_model=Token)
def login(
    form_data: OAuth2PasswordRequestForm = Depends(),
    db: Session = Depends(get_db)
):
    user = authenticate_user(db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token = create_access_token(
        data={"sub": user.username},
        expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    )
    return {"access_token": access_token, "token_type": "bearer"}


@app.get("/users/me")
def read_users_me(current_user: User = Depends(get_current_user)):
    return {"username": current_user.username, "email": current_user.email}


# ── Document Upload & Ingestion ───────────────────────────────

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    ALLOWED_EXTS = {".pdf", ".docx", ".xlsx", ".xls", ".txt"}
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in ALLOWED_EXTS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: PDF, DOCX, XLSX, TXT"
        )

    # FIX: Use pathlib for Windows-compatible paths
    session_id = str(uuid.uuid4())
    safe_name = f"{session_id}_{file.filename}"
    save_path = UPLOAD_DIR / safe_name

    with open(save_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Build vectorstore (with Ollama error handling)
    try:
        vectorstore, num_chunks = build_vectorstore(session_id, save_path)
    except HTTPException as e:
        save_path.unlink(missing_ok=True)
        raise e
    except Exception as e:
        save_path.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")

    # Save document record
    doc = Document(
        filename=file.filename,
        file_path=str(save_path),
        owner_id=current_user.id,
        session_id=session_id
    )
    db.add(doc)
    db.commit()

    return {
        "message": "Document uploaded and indexed successfully",
        "session_id": session_id,
        "filename": file.filename,
        "chunks_indexed": num_chunks
    }


# ── Q&A / RAG ─────────────────────────────────────────────────

@app.post("/ask", response_model=AnswerResponse)
async def ask_question(
    request: QuestionRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if not check_ollama():
        raise HTTPException(
            status_code=503,
            detail="Ollama is not running. Start it with: ollama serve"
        )

    vectorstore = get_vectorstore(request.session_id)
    if vectorstore is None:
        raise HTTPException(
            status_code=404,
            detail="Session not found. Please upload a document first."
        )

    try:
        llm = OllamaLLM(model=OLLAMA_MODEL, temperature=0.1)

        prompt_template = """You are a helpful research assistant. 
Use the following context to answer the question. 
If you don't know the answer from the context, say so clearly.

Context:
{context}

Question: {question}

Answer:"""

        prompt = PromptTemplate(
            input_variables=["context", "question"],
            template=prompt_template
        )

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
            chain_type_kwargs={"prompt": prompt},
            return_source_documents=True
        )

        result = qa_chain.invoke({"query": request.question})
        answer = result.get("result", "No answer generated.")
        sources = list({
            doc.metadata.get("source", "Unknown")
            for doc in result.get("source_documents", [])
        })

        # Save to chat history
        history = ChatHistory(
            session_id=request.session_id,
            question=request.question,
            answer=answer
        )
        db.add(history)
        db.commit()

        return AnswerResponse(
            answer=answer,
            session_id=request.session_id,
            sources=sources
        )

    except httpx.ConnectError:
        raise HTTPException(
            status_code=503,
            detail="Ollama disconnected. Make sure 'ollama serve' is running."
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Inference error: {str(e)}")


# ── Streaming Q&A (Bonus Feature) ────────────────────────────

@app.post("/ask/stream")
async def ask_stream(
    request: QuestionRequest,
    current_user: User = Depends(get_current_user)
):
    """Streams answer tokens as they are generated — better UX for slow hardware."""
    if not check_ollama():
        raise HTTPException(status_code=503, detail="Ollama is not running.")

    vectorstore = get_vectorstore(request.session_id)
    if vectorstore is None:
        raise HTTPException(status_code=404, detail="Session not found.")

    docs = vectorstore.similarity_search(request.question, k=3)
    context = "\n\n".join([doc.page_content for doc in docs])

    prompt = f"""You are a helpful research assistant.

Context:
{context}

Question: {request.question}

Answer:"""

    llm = OllamaLLM(model=OLLAMA_MODEL, temperature=0.1)

    async def token_stream():
        try:
            async for chunk in llm.astream(prompt):
                yield chunk
        except httpx.ConnectError:
            yield "\n[Error: Ollama disconnected]"

    return StreamingResponse(token_stream(), media_type="text/plain")


# ── Chat History ──────────────────────────────────────────────

@app.get("/history/{session_id}")
def get_chat_history(
    session_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    history = (
        db.query(ChatHistory)
        .filter(ChatHistory.session_id == session_id)
        .order_by(ChatHistory.created_at)
        .all()
    )
    return [
        {"question": h.question, "answer": h.answer, "time": h.created_at}
        for h in history
    ]


@app.get("/documents")
def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    docs = db.query(Document).filter(Document.owner_id == current_user.id).all()
    return [
        {
            "filename": d.filename,
            "session_id": d.session_id,
            "uploaded_at": d.uploaded_at
        }
        for d in docs
    ]


@app.delete("/documents/{session_id}")
def delete_document(
    session_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(Document).filter(
        Document.session_id == session_id,
        Document.owner_id == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove file and vectorstore cache
    file_path = Path(doc.file_path)
    if file_path.exists():
        file_path.unlink()
    _vectorstore_cache.pop(session_id, None)

    db.delete(doc)
    db.commit()
    return {"message": "Document deleted successfully"}


# ─────────────────────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
